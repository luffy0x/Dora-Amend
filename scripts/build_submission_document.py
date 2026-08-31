from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "document" / "Dora-Amend-参赛设计文档.docx"
IMAGES = ROOT / "docs" / "submission-assets"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(23, 35, 59)
MUTED = RGBColor(91, 105, 130)
LIGHT = "EAF1FF"
LATIN_FONT = "Arial"
CJK_FONT = "Arial Unicode MS"


def set_ooxml_fonts(r_fonts):
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT)
    r_fonts.set(qn("w:cs"), LATIN_FONT)


def set_run_font(run, size=11, bold=False, color=DARK):
    run.font.name = LATIN_FONT
    set_ooxml_fonts(run._element.get_or_add_rPr().rFonts)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    set_ooxml_fonts(fonts)
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    run.append(properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_heading(doc, text):
    paragraph = doc.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text), size=16, bold=True, color=BLUE)
    return paragraph


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    set_run_font(paragraph.add_run(text), size=9, color=MUTED)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Dora Amend  |  ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    set_ooxml_fonts(normal._element.get_or_add_rPr().rFonts)
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK

    for style_name, size, before, after in [
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        set_ooxml_fonts(style._element.get_or_add_rPr().rFonts)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def new_page(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    set_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "2026 帆软 AI 产品体验设计挑战赛  |  Dora 赛道 命题 3"
    set_run_font(header.runs[0], size=9, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(12)
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("可信回答体验方案"), size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("Dora Amend"), size=28, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("让每个结论都看得懂、信得过、能追溯，也能被纠正"), size=14, color=MUTED)

    meta = doc.add_table(rows=1, cols=2)
    meta.autofit = False
    meta.columns[0].width = Inches(3.5)
    meta.columns[1].width = Inches(3.5)
    for cell in meta.rows[0].cells:
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = 1
    left = meta.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(left.add_run("参赛者：王志豪"), size=10.5, bold=True)
    right = meta.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(right.add_run("交付形式：React 高保真原型"), size=10.5, bold=True)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(16)
    lead.paragraph_format.space_after = Pt(14)
    set_run_font(
        lead.add_run("把可信体验从“展示更多日志”推进到“围绕结论核验、确认影响、修正并比较”。"),
        size=12,
        bold=True,
        color=BLUE,
    )

    doc.add_picture(str(IMAGES / "01-trust-overview.png"), width=Inches(7.0))
    add_caption(doc, "图 1  可信摘要与结论级证据，来源、范围、筛选和事实推断关系集中呈现")

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_before = Pt(6)
    links.paragraph_format.space_after = Pt(0)
    add_hyperlink(links, "在线 Demo", "https://luffy0x.github.io/Dora-Amend/")
    set_run_font(links.add_run("  |  "), size=10, color=MUTED)
    add_hyperlink(links, "代码仓库", "https://github.com/luffy0x/Dora-Amend")

    new_page(doc)
    add_heading(doc, "问题洞察")
    add_body(
        doc,
        "Dora 已具备思考过程、工具步骤、SQL 明细和最佳实践引用，但这些信息分散在会话与后台监控中。普通业务用户看到经营结论时，仍难以快速确认数据来源、更新时间、筛选条件，以及哪些是事实或推断。结论有误时，还需要重新发起对话或寻找后台入口。问题不是日志不够多，而是证据没有围绕用户正在判断的结论组织。",
    )
    add_heading(doc, "设计方案")
    add_body(
        doc,
        "Dora Amend 将回答重组为可核验、可修订、可对比的结论对象。回答顶部用可信摘要集中呈现数据来源、更新时间、分析范围、关键口径和异常提醒，并区分数据事实、Agent 推断与行动建议。用户点击任一结论，只查看与它直接相关的事实、计算、筛选条件、查询记录和最佳实践引用，不必翻阅完整日志。",
    )
    add_body(
        doc,
        "当用户发现“新店与成熟门店被直接比较”等问题，可在当前结论旁发起纠正。自然语言先转换为明确的结构化条件，执行前展示哪些事实、推断和建议会受影响，确认后仅局部重跑。系统保留原回答，并用版本对比突出利润降幅、原因排序和行动建议为何变化。",
    )
    doc.add_picture(str(IMAGES / "02-impact-preview.png"), width=Inches(7.0))
    add_caption(doc, "图 2  影响预览先说明会重算什么、保持什么，再由用户确认执行")

    new_page(doc)
    add_heading(doc, "创新与价值")
    add_body(
        doc,
        "方案不使用难以解释的可信度百分比，也不展示模型隐藏思维过程，而是提供可审计的输入、系统动作和工具结果。它把可信体验从只读解释推进为发现问题、确认影响、修正结论的闭环，让用户知道结论为什么变化，也避免基于错误口径直接采取行动。",
    )
    doc.add_picture(str(IMAGES / "03-version-comparison.png"), width=Inches(7.0))
    add_caption(doc, "图 3  修订后保留 v1，并明确展示数值、原因排序与行动建议的变化")
    add_heading(doc, "验证与边界")
    add_body(
        doc,
        "当前交付为 React 高保真交互原型，使用虚构 Mock 数据模拟完整闭环，并覆盖数据过期、来源不可用、权限不足和修正歧义状态。原型不连接真实数据库、模型或 Text-to-SQL 服务。类型检查、生产构建、响应式和线上黄金路径已完成验证；真实用户测试尚未完成，不声明未经验证的效果。",
    )

    for current_section in doc.sections:
        current_section.page_width = Inches(8.5)
        current_section.page_height = Inches(11)
        current_section.top_margin = Inches(0.7)
        current_section.bottom_margin = Inches(0.7)
        current_section.left_margin = Inches(0.75)
        current_section.right_margin = Inches(0.75)
        current_section.header_distance = Inches(0.35)
        current_section.footer_distance = Inches(0.35)

    doc.core_properties.title = "Dora Amend 参赛设计文档"
    doc.core_properties.subject = "2026 帆软 AI 产品体验设计挑战赛 Dora 赛道命题 3"
    doc.core_properties.author = "王志豪"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
